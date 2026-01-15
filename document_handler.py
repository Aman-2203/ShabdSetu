import time
import logging
from typing import List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class DocumentHandler:
    """Handles document I/O"""
    
    @staticmethod
    def read_docx(file_path: str) -> str:
        """Read Word document"""
        try:
            doc = Document(file_path)
            content = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
            return '\n\n'.join(content)
        except Exception as e:
            raise Exception(f"Error reading document: {e}")
    
    
    @staticmethod
    def create_formatted_document(chunks: List[str], output_file: str, 
                                language: str, doc_type: str = "Processed"):
        """Create formatted Word document with improved formatting"""
        try:
            if language.lower() == 'gujarati':
                primary_font = 'Noto Serif Gujarati'
                sanskrit_font = 'Noto Serif Devanagari'
            else:
                primary_font = 'Noto Serif Devanagari'
                sanskrit_font = 'Noto Serif Devanagari'
        
            new_doc = Document()
            
            style = new_doc.styles['Normal']
            font = style.font
            font.name = primary_font
            font.size = Pt(12)
            
            paragraph_format = style.paragraph_format
            paragraph_format.first_line_indent = Inches(0.5)
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            title = new_doc.add_heading(f'{doc_type} {language.title()} Document', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.name = primary_font
                run.font.size = Pt(18)
            
            subtitle = new_doc.add_paragraph()
            subtitle.add_run(f"{doc_type} Version").italic = True
            subtitle.add_run(f"\nProcessed on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            subtitle.add_run(f"\nLanguage: {language.title()}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in subtitle.runs:
                run.font.name = primary_font
                run.font.size = Pt(11)
            
            new_doc.add_page_break()
            
            def is_heading(text):
                text = text.strip()
                return text.startswith('**') and text.endswith('**')
            
            def extract_heading_text(text):
                text = text.strip()
                if text.startswith('**') and text.endswith('**'):
                    return text[2:-2].strip()
                return text
            
            def has_sanskrit_formatting(text):
                return '<' in text and '>' in text
            
            def process_sanskrit_text(para, text):
                if not has_sanskrit_formatting(text):
                    run = para.add_run(text)
                    run.font.name = primary_font
                    run.font.size = Pt(12)
                    return
                
                current_pos = 0
                while current_pos < len(text):
                    start_marker = text.find('<', current_pos)
                    if start_marker == -1:
                        remaining_text = text[current_pos:]
                        if remaining_text.strip():
                            run = para.add_run(remaining_text)
                            run.font.name = primary_font
                            run.font.size = Pt(12)
                        break
                    before_text = text[current_pos:start_marker]
                    if before_text.strip():
                        run = para.add_run(before_text)
                        run.font.name = primary_font
                        run.font.size = Pt(12)
                    end_marker = text.find('>', start_marker)
                    if end_marker == -1:
                        remaining_text = text[start_marker+1:]
                        run = para.add_run(remaining_text)
                        run.font.name = primary_font
                        run.font.size = Pt(12)
                        break
                    sanskrit_text = text[start_marker+1:end_marker]
                    if sanskrit_text.strip():
                        sanskrit_run = para.add_run(sanskrit_text)
                        sanskrit_run.font.name = sanskrit_font
                        sanskrit_run.font.size = Pt(12)
                        sanskrit_run.italic = True
                    current_pos = end_marker + 1
            
            # Process chunks
            for chunk_text in chunks:
                if chunk_text and chunk_text.strip():
                    paragraphs = chunk_text.split('\n\n')
                    
                    for para_text in paragraphs:
                        para_text = para_text.strip()
                        if para_text:
                            if is_heading(para_text):
                                heading_text = extract_heading_text(para_text)
                                heading = new_doc.add_heading(heading_text, level=2)
                                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in heading.runs:
                                    run.font.name = primary_font
                                    run.font.size = Pt(14)
                                    run.font.bold = True
                            else:
                                # Split into individual lines and process each
                                lines = para_text.split('\n')
                                
                                for line in lines:
                                    line = line.strip()
                                    if not line:
                                        continue
                                    
                                    # Determine alignment based on character count (excluding formatting markers)
                                    clean_text = line.replace('<', '').replace('>', '')
                                    char_count = len(clean_text.strip())
                                    alignment = WD_ALIGN_PARAGRAPH.LEFT if char_count <= 20 else WD_ALIGN_PARAGRAPH.JUSTIFY
                                    
                                    # Colon-bold formatting for non-headings
                                    if ':' in line and not has_sanskrit_formatting(line.split(':', 1)[0]):
                                        para = new_doc.add_paragraph()
                                        para.paragraph_format.first_line_indent = Inches(0.3)
                                        para.paragraph_format.alignment = alignment
                                        para.paragraph_format.space_before = Pt(0.75)
                                        para.paragraph_format.space_after = Pt(0.75)
                                        parts = line.split(':', 1)
                                        if len(parts) == 2:
                                            bold_run = para.add_run(parts[0] + ':')
                                            bold_run.bold = False
                                            bold_run.font.name = primary_font
                                            bold_run.font.size = Pt(12)
                                            remaining_text = ' ' + parts[1]
                                            process_sanskrit_text(para, remaining_text)
                                        else:
                                            process_sanskrit_text(para, line)
                                    else:
                                        para = new_doc.add_paragraph()
                                        para.paragraph_format.first_line_indent = Inches(0.3)
                                        para.paragraph_format.alignment = alignment
                                        para.paragraph_format.space_before = Pt(0.75)
                                        para.paragraph_format.space_after = Pt(0.75)
                                        process_sanskrit_text(para, line)
            
            new_doc.save(output_file)
            return True
            
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            return False

    @staticmethod
    def save_raw_docx(text: str, output_path: str):
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        paragraph_format = style.paragraph_format
        paragraph_format.first_line_indent = Inches(0.5)
        # Use left alignment for raw OCR output
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        title = doc.add_heading(f'{"OCR"}  Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(18)
            
        subtitle = doc.add_paragraph()
        subtitle.add_run("OCR").italic = True
        subtitle.add_run(f"\nProcessed on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        
        doc.add_page_break()


        for paragraph in text.split('\n\n'):
            if paragraph.strip():
                para = doc.add_paragraph(paragraph.strip())
                # Ensure each paragraph is left-aligned explicitly
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        

        doc.save(output_path)