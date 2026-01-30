import os
import math
import logging
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

try:
    import fitz  # PyMuPDF
    from docx import Document
except ImportError as e:
    print(f"Missing required package: {e}")
    print("\nPlease install: pip install PyMuPDF python-docx")
    import sys
    sys.exit(1)

logger = logging.getLogger(__name__)

# Email configuration
GMAIL_ADDRESS = os.getenv('SENDER_EMAIL', '')
GMAIL_APP_PASSWORD = os.getenv('SENDER_PASSWORD', '')

# Trial limits
TRIAL_PAGE_LIMIT = 3
TRIAL_CHAR_LIMIT = 10000
CHARS_PER_PAGE = 3333  # 10,000 chars / 3 pages = ~3,333 chars per page

# Maximum page limit for PDFs (for both trial and paid users)
MAX_PDF_PAGES = 200


def get_pdf_page_count(file_path):
    """
    Get the number of pages in a PDF file.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        int: Number of pages in the PDF
        
    Raises:
        Exception: If file cannot be read or is not a valid PDF or exceeds max page limit
    """
    try:
        pdf_document = fitz.open(file_path)
        page_count = pdf_document.page_count
        pdf_document.close()
        logger.info(f"PDF has {page_count} pages: {file_path}")
        
        # Check if PDF exceeds maximum page limit
        if page_count > MAX_PDF_PAGES:
            raise Exception(
                f"PDF exceeds maximum page limit. Your PDF has {page_count} pages, "
                f"but the maximum allowed is {MAX_PDF_PAGES} pages. "
                f"Please split your document into smaller files."
            )
        
        return page_count
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
        if "exceeds maximum page limit" in str(e):
            raise  # Re-raise the max page limit error as-is
        raise Exception(f"Unable to read PDF file: {str(e)}")


def get_docx_character_count(file_path):
    """
    Get the total character count in a Word document.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        int: Total character count (excluding spaces)
        
    Raises:
        Exception: If file cannot be read or is not a valid DOCX
    """
    try:
        doc = Document(file_path)
        total_chars = 0
        
        # Count characters in all paragraphs
        for paragraph in doc.paragraphs:
            # Count characters excluding spaces
            text = paragraph.text
            total_chars += len(text.replace(' ', ''))
        
        logger.info(f"DOCX has {total_chars} characters: {file_path}")
        return total_chars
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        raise Exception(f"Unable to read Word document: {str(e)}")


def calculate_page_usage(file_path, file_extension):
    """
    Calculate page usage based on file type.
    
    Args:
        file_path (str): Path to the file
        file_extension (str): File extension (with or without dot)
        
    Returns:
        dict: {
            'page_usage': float,  # Page units to deduct from trial
            'actual_pages': int or None,  # Actual page count for PDFs
            'char_count': int or None,  # Character count for DOCX
            'file_type': str  # 'pdf' or 'docx'
        }
        
    Raises:
        Exception: If file type is unsupported or file cannot be read
    """
    # Normalize extension
    extension = file_extension.lower().strip('.')
    
    if extension == 'pdf':
        page_count = get_pdf_page_count(file_path)
        return {
            'page_usage': float(page_count),
            'actual_pages': page_count,
            'char_count': None,
            'file_type': 'pdf'
        }
    elif extension in ['docx', 'doc']:
        char_count = get_docx_character_count(file_path)
        # Convert characters to page equivalents (round up)
        page_usage = math.ceil(char_count / CHARS_PER_PAGE * 100) / 100  # Round to 2 decimals
        return {
            'page_usage': page_usage,
            'actual_pages': None,
            'char_count': char_count,
            'file_type': 'docx'
        }
    else:
        raise Exception(f"Unsupported file type: {extension}")


def validate_trial_limits(page_usage_info, remaining_pages):
    """
    Validate if the document can be processed within trial limits.
    
    Args:
        page_usage_info (dict): Result from calculate_page_usage()
        remaining_pages (float): Remaining trial pages for the user
        
    Returns:
        dict: {
            'valid': bool,
            'message': str,  # Error message if not valid
            'page_usage': float
        }
    """
    page_usage = page_usage_info['page_usage']
    file_type = page_usage_info['file_type']
    
    # Check if user has enough pages remaining
    if page_usage > remaining_pages:
        if file_type == 'pdf':
            actual_pages = page_usage_info['actual_pages']
            message = (
                f"Your PDF document has {actual_pages} page{'s' if actual_pages > 1 else ''}, "
                f"but you only have {remaining_pages:.1f} page{'s' if remaining_pages != 1 else ''} remaining in your trial. "
                f"Please upload a smaller document or upgrade your account."
            )
        else:  # docx
            char_count = page_usage_info['char_count']
            message = (
                f"Your Word document has {char_count:,} characters (≈{page_usage:.1f} pages), "
                f"but you only have {remaining_pages:.1f} page{'s' if remaining_pages != 1 else ''} remaining in your trial. "
                f"Maximum allowed: {TRIAL_CHAR_LIMIT:,} characters (≈{TRIAL_PAGE_LIMIT} pages). "
                f"Please upload a smaller document or upgrade your account."
            )
        
        return {
            'valid': False,
            'message': message,
            'page_usage': page_usage
        }
    
    # Check absolute limits (3 pages or 10k chars)
    if file_type == 'pdf' and page_usage_info['actual_pages'] > TRIAL_PAGE_LIMIT:
        message = (
            f"Your PDF document has {page_usage_info['actual_pages']} pages, "
            f"which exceeds the trial limit of {TRIAL_PAGE_LIMIT} pages per document."
        )
        return {
            'valid': False,
            'message': message,
            'page_usage': page_usage
        }
    elif file_type == 'docx' and page_usage_info['char_count'] > TRIAL_CHAR_LIMIT:
        message = (
            f"Your Word document has {page_usage_info['char_count']:,} characters, "
            f"which exceeds the trial limit of {TRIAL_CHAR_LIMIT:,} characters per document "
            f"(≈{TRIAL_PAGE_LIMIT} pages)."
        )
        return {
            'valid': False,
            'message': message,
            'page_usage': page_usage
        }
    
    return {
        'valid': True,
        'message': 'Document is within trial limits',
        'page_usage': page_usage
    }


def get_docx_word_count(file_path):
    """
    Get the total word count in a Word document.
    For paid users, this is used to estimate page count (500 words per page).
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        int: Total word count
        
    Raises:
        Exception: If file cannot be read or is not a valid DOCX
    """
    try:
        doc = Document(file_path)
        total_words = 0
        
        # Count words in all paragraphs
        for paragraph in doc.paragraphs:
            # Split by whitespace and count non-empty strings
            words = paragraph.text.split()
            total_words += len([w for w in words if w.strip()])
        
        logger.info(f"DOCX has {total_words} words: {file_path}")
        return total_words
    except Exception as e:
        logger.error(f"Error reading DOCX for word count: {e}")
        raise Exception(f"Unable to read Word document: {str(e)}")


def calculate_pages_from_words(word_count, words_per_page=550):
    """
    Calculate estimated page count from word count.
    Used for paid users processing DOCX files.
    
    Args:
        word_count (int): Total word count
        words_per_page (int): Words per page (default: 550)
        
    Returns:
        int: Estimated page count (rounded up)
    """
    if word_count <= 0:
        return 0
    return math.ceil(word_count / words_per_page)


def send_document_email(recipient_email, document_path, job_id):
    """
    Send processed document to user via email.
    
    Args:
        recipient_email (str): Recipient's email address
        document_path (str): Path to the processed document file
        job_id (str): Job ID for reference
        
    Returns:
        bool: True if email sent successfully, False otherwise
    """
    try:
        # Check if file exists
        if not os.path.exists(document_path):
            logger.error(f"Document not found: {document_path}")
            return False
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = GMAIL_ADDRESS
        msg['To'] = recipient_email
        msg['Subject'] = f'Your Processed Document - Job {job_id[:8]}'
        
        # Email body
        filename = os.path.basename(document_path)
        body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <h2 style="color: #4CAF50;">Document Processing Complete!</h2>
                <p>Your document has been successfully processed and is ready for download.</p>
                
                <div style="background-color: #f5f5f5; padding: 15px; border-radius: 5px; margin: 20px 0;">
                    <p><strong>Job ID:</strong> {job_id}</p>
                    <p><strong>Document:</strong> {filename}</p>
                </div>
                
                <p>Please find your processed document attached to this email.</p>
                
                <p style="margin-top: 30px; padding-top: 20px; border-top: 1px solid #ddd; color: #666; font-size: 12px;">
                    This is an automated email from Shabd Setu Document Processing Service.<br>
                    If you didn't request this document, please ignore this email.
                </p>
            </div>
        </body>
        </html>
        """
        
        msg.attach(MIMEText(body, 'html'))
        
        # Attach the document
        with open(document_path, 'rb') as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
        
        encoders.encode_base64(part)
        part.add_header(
            'Content-Disposition',
            f'attachment; filename= {filename}'
        )
        
        msg.attach(part)
        
        # Send email
        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(GMAIL_ADDRESS, GMAIL_APP_PASSWORD)
            server.send_message(msg)
        
        logger.info(f"Document email sent successfully to {recipient_email}")
        return True
        
    except Exception as e:
        logger.error(f"Failed to send document email: {e}")
        return False
